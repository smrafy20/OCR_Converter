import cv2
import numpy as np
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from docx import Document
import os
import time

# Set Tesseract path
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def process_pdf(pdf_path, output_word_path):
    try:
        # Create a temporary folder for images
        temp_folder = "temp_images"
        if not os.path.exists(temp_folder):
            os.makedirs(temp_folder)
        
        # Convert PDF to images
        print("Converting PDF to images...")
        images = convert_from_path(pdf_path)
        
        # Create a new Word document
        doc = Document()
        
        # Process each page
        for i, image in enumerate(images):
            # Save the image temporarily
            image_path = f"{temp_folder}/page_{i+1}.jpg"
            image.save(image_path, "JPEG")
            
            print(f"Processing page {i+1}...")
            
            # Open image with OpenCV
            img = cv2.imread(image_path)
            
            # Preprocess for better OCR results
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            blur = cv2.GaussianBlur(gray, (5,5), 0)
            _, threshold = cv2.threshold(blur, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # Save processed image
            processed_image_path = f"{temp_folder}/processed_{i+1}.jpg"
            cv2.imwrite(processed_image_path, threshold)
            
            # Use multi-language approach for mixed content
            try:
                # Try with Bengali+English combined
                text = pytesseract.image_to_string(
                    Image.open(processed_image_path), 
                    lang='ben+eng'  # Use both languages together
                )
                
                # If text is very short or empty, try alternate approaches
                if len(text.strip()) < 10:
                    # Try English only
                    eng_text = pytesseract.image_to_string(
                        Image.open(processed_image_path), 
                        lang='eng'
                    )
                    
                    # Try Bengali only
                    ben_text = pytesseract.image_to_string(
                        Image.open(processed_image_path), 
                        lang='ben'
                    )
                    
                    # Use the result with more content
                    if len(eng_text) > len(ben_text):
                        text = eng_text
                        print(f"Used English OCR for page {i+1}")
                    else:
                        text = ben_text
                        print(f"Used Bengali OCR for page {i+1}")
                else:
                    print(f"Used combined Bengali+English OCR for page {i+1}")
                
                # Add text to document
                doc.add_paragraph(text)
                
            except Exception as e:
                print(f"OCR error on page {i+1}: {e}")
                # Try fallback to individual languages
                try:
                    eng_text = pytesseract.image_to_string(
                        Image.open(processed_image_path), 
                        lang='eng'
                    )
                    doc.add_paragraph(eng_text)
                    print(f"Fallback to English OCR for page {i+1}")
                except Exception as eng_e:
                    print(f"English OCR error: {eng_e}")
                    try:
                        ben_text = pytesseract.image_to_string(
                            Image.open(processed_image_path), 
                            lang='ben'
                        )
                        doc.add_paragraph(ben_text)
                        print(f"Fallback to Bengali OCR for page {i+1}")
                    except Exception as ben_e:
                        print(f"Bengali OCR error: {ben_e}")
                        doc.add_paragraph(f"[OCR failed for page {i+1}]")
        
        # Try to save with a unique filename
        timestamp = int(time.time())
        new_output_path = f"bangla_english_text_{timestamp}.docx"
        
        try:
            doc.save(new_output_path)
            print(f"Word document saved to {new_output_path}")
        except Exception as save_error:
            print(f"Error saving to default location: {save_error}")
            # Try saving to Desktop as fallback
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            alt_output_path = os.path.join(desktop_path, new_output_path)
            try:
                doc.save(alt_output_path)
                print(f"Word document saved to alternative location: {alt_output_path}")
            except Exception as alt_save_error:
                print(f"Error saving to alternative location: {alt_save_error}")
        
        # Clean up temporary files
        for file in os.listdir(temp_folder):
            os.remove(os.path.join(temp_folder, file))
        os.rmdir(temp_folder)
        
    except Exception as e:
        print(f"Error: {e}")

# Run the conversion
process_pdf("test1.pdf", "bangla_english_text.docx")